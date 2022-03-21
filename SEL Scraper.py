# 2021-03-16
#webdriver from google chrome:    https://www.youtube.com/watch?v=b5jt2bhSeXs
from typing import Set
from selenium.webdriver.common.by import By
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
# from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
# from selenium.webdriver.common.action_chains import ActionChains
import pandas as pd, numpy as np
import datetime, time, random, glob, os, shutil, re
from docx import Document

def main():
	#inputs
	path = r'C:\Users\yongj\Downloads\2021-03-16 SEL Scraper\Links.xlsx'
	path_storage = r'C:\Users\yongj\Downloads\3. Business News\To Read'
	headless_window = True
	# headless_window = False
	idx_link = 0


	#Part One: Get links from excel
	PartOne = ObtainLinks(path)
	list_links = PartOne.ReadExcel()


	#setting up the options for the chromedriver for Part Two
	options = Options()
	if headless_window == True:
		options.add_argument('--headless')
		options.add_argument("window-size=1920,1080") 
	#defining the window size of the headless browser avoids errors with finding location and date for each job posting
	path_chromedriver = r'C:\Program Files (x86)\chromedriver.exe'
	driver = webdriver.Chrome(executable_path=path_chromedriver, options=options)

	#Part Two: Webrscrape
	PartTwo = Webscrape(path_storage)
	while idx_link < len(list_links):
		link = list_links[idx_link]
		print(f'{idx_link+1}/{len(list_links)})\t{link}')
		article_type = PartTwo.Article_type(link)
		# print(article_type)
		PartTwo.AccessWebpage(driver, link, article_type)
		print('\n')	
		idx_link+=1


	#Part Three: Moving Word Docs to the "To Read Folder" & updating excel
	#Refer to python saved in Superseded to move all word docs to 'To Read' folder and update the excel
	import sys
	sys.path.insert(1, 'C:/Users/yongj/Documents/Coding/Python/Web Scraping/0. Superseded')
	from Update_Excel_Move_Word import UpdateExcel_MoveWord
	UpdateExcel_MoveWord()



class ObtainLinks():
	def __init__(self, path):
		self.path = path

	def ReadExcel(self):
		# print(self.path)
		df = pd.read_excel(self.path)
		# print(df['No Seekingalpha Links'].to_list()[:10])
		self.list_links = df['No Seekingalpha Links'].to_list()
		return self.list_links

class SEL_Scrapers():
	def __init__(self, driver, link):
		self.driver = driver
		self.link = link

	# comment below is for reference for the pieces of info required
	# def Create_docx(self, url, title, date, author, summary, content, table_data, img_tags):
	def Create_docx(self):    
	    url = self.link
	    title = self.title
	    date = self.date
	    author = self.author
	    summary = self.summary
	    content = self.content
	    table_data = self.table_data
	    img_tags = self.img_tags

	    document = Document()
	        
	    document.add_heading(url)

	    #Removing characters that cannot be used in the title
	    title = title.replace('/', " ")
	    title = title.replace('"', '')
	    title = title.replace(':', ' - ') 
	    title = title.replace('?', ' - ')
	    title = title.replace('*', ' - ')
	    title = title.replace('|', ' - ')

	    #add in content to the word docx
	    document.add_paragraph(title)
	    document.add_paragraph(str(date))
	    document.add_paragraph(author)
	    for a in summary:
	        document.add_paragraph(a.text)
	    row=1
	    for p in content:
	    # 	document.add_paragraph(f'Line {row}: \n{p.text}')
	    # 	row+=1
	    	if len(p.text.split()) >3:
	        	document.add_paragraph(p.text)

	    #Create a table if table data is passed in
	    if table_data:
	        # rowNum = len(table_data)
	        colNum = len(table_data[0])
	        #initialize table with a blank row
	        Table = document.add_table(rows=1,cols=colNum)
	        # # table styles link:  https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html#table-styles-in-default-template
	        Table.style = "Table Grid"
	        for row in table_data:
	            row_Cells = Table.add_row().cells 
	            for cell in range(0,colNum):
	                row_Cells[cell].text = row[cell]    


	    #Append Images to word doc
	    links =[]

	    for img in img_tags:
	        try:
	            img_width = img['data-width']
	            links.append(img['src'])
	        except KeyError:
	            continue
	    # print(links)
	    list_imgtypes = ['.jpg','.png']
	    num = 1
	    for link in links:
	        #Get image type
	        image_type = link[-4:]
	        if image_type in list_imgtypes:
	            image_type=image_type
	        else:
	            image_type='.jpeg'
	        
	        try:
	            filename = "Image"+str(num)
	            imagefile = open(filename+image_type, "wb")
	            imagefile.write(requests.get(link).content)
	            imagefile.close()
	            # print('Created image')
	            document.add_paragraph(filename+image_type)
	            document.add_picture(filename+image_type,width=docx.shared.Inches(6))
	            # document.add_picture(requests.get(link).content,width=Inches(4))
	            os.remove(filename+image_type)
	            # print(f'added {filename}')

	        except:
	            # print('Then Failed to add')
	            filename = "Image"+str(num)
	            document.add_paragraph(filename+image_type + " COULD NOT BE ADDED DUE TO SOME ERROR.\n"+link)
	        num+=1

	    wordname = str(date)+" " + title+".docx"
	    document.save(wordname)
	    # path_toREAD = r"C:\Users\yongj\Downloads\3. Business News\To Read"+'\\'+wordname
	    # document.save(path_toREAD)
	def Reset_PgINFO(self):
		self.url = ''
		self.title =''
		self.date =''
		self.author=''
		self.summary= []
		self.content=[]
		self.table_data=[]
		self.img_tags=[]


	def oilprice(self):
		# def Create_docx(self, url, title, date, author, summary, content, table_data, img_tags):
		print('Webscraping link from: "Oilprice.com"')
		driver = self.driver

		self.title = driver.find_element_by_xpath('//*[@id="pagecontent"]/div[3]/div/div[1]/div[2]/div[2]/h1').text
		self.author =  driver.find_element_by_xpath('//*[@id="pagecontent"]/div[3]/div/div[1]/div[2]/div[2]/span/a').text
		self.date =  driver.find_element_by_xpath('//*[@id="pagecontent"]/div[3]/div/div[1]/div[2]/div[2]/span').text
		self.date = Webscrape.FormatDate(self, self.date)
		self.summary =  []
		self.content =  driver.find_elements_by_xpath('//*[@id="news-content"]')
		self.table_data =  []
		self.img_tags =  []

		print(self.title)
		# for p in self.content:
		# 	print(p.text)

	def boereport(self):
		# def Create_docx(self, url, title, date, author, summary, content, table_data, img_tags):
		print('Webscraping link from: "boereport.com"')
		driver = self.driver

		self.title = driver.find_element_by_xpath('/html/body/div[1]/div/div/main/article/header/h1').text
		self.author =  'boereport'
		self.date =  driver.find_element_by_xpath('/html/body/div[1]/div/div/main/article/header/p/time[1]').text
		self.date = Webscrape.FormatDate(self, self.date)
		self.summary =  []
		self.content =  driver.find_elements_by_xpath('/html/body/div[1]/div/div/main/article/div')
		self.table_data =  []
		self.img_tags =  []

	def cnbc(self):
		# def Create_docx(self, url, title, date, author, summary, content, table_data, img_tags):
		print('Webscraping link from: "cnbc.com"')
		driver = self.driver

		self.title = driver.find_element_by_xpath('//*[@id="main-article-header"]/div/div[1]/div[1]/h1').text
		self.author =  driver.find_element_by_xpath('//*[@id="main-article-header"]/div/div[2]/div[1]/div/div/div/div/a[1]').text
		self.date =  driver.find_element_by_xpath('//*[@id="main-article-header"]/div/div[1]/div[2]/time[1]').text
		self.date = Webscrape.FormatDate(self, self.date)
		self.summary =  driver.find_elements_by_xpath('//*[@id="RegularArticle-KeyPoints-4"]/div/div[2]/div/div/ul')
		self.content =  driver.find_elements_by_xpath('//*[@id="RegularArticle-ArticleBody-5"]/div[2]')
		self.table_data =  []
		self.img_tags =  []

		print(self.title,'\n')
		# for p in self.content:
		# 	print(p.text)

	def nbc(self):
		# def Create_docx(self, url, title, date, author, summary, content, table_data, img_tags):
		print('Webscraping link from: "nbc.com"')
		driver = self.driver

		self.title = driver.find_element_by_xpath('//*[@id="main-article-header"]/div/div[1]/div[1]/h1').text
		self.author =  driver.find_element_by_xpath('//*[@id="main-article-header"]/div/div[2]/div[1]/div/div/div/div/a[1]').text
		self.date =  driver.find_element_by_xpath('//*[@id="main-article-header"]/div/div[1]/div[2]/time[1]').text
		self.date = Webscrape.FormatDate(self, self.date)
		self.summary =  driver.find_elements_by_xpath('//*[@id="RegularArticle-KeyPoints-4"]/div/div[2]/div/div/ul')
		self.content =  driver.find_elements_by_xpath('//*[@id="RegularArticle-ArticleBody-5"]/div[2]')
		self.table_data =  []
		self.img_tags =  []

		print(self.title,'\n')
		# for p in self.content:
		# 	print(p.text)

	def AB(self):
		# def Create_docx(self, url, title, date, author, summary, content, table_data, img_tags):
		print('Webscraping link from: "ab.ca"')
		driver = self.driver

		self.title = driver.find_element_by_xpath('//*[@id="top"]/header/div[4]/div[2]/div/h1').text
		self.author =  'AB Government'
		self.date =  driver.find_element_by_xpath('//*[@id="top"]/header/div[4]/div[2]/div/div/time').text
		self.date = Webscrape.FormatDate(self, self.date)
		self.summary =  driver.find_elements_by_xpath('//*[@id="top"]/header/div[4]/div[2]/div')
		self.content =  driver.find_elements_by_xpath('//*[@id="main"]/div[2]/div')
		self.table_data =  []
		self.img_tags =  []

		print(self.title,'\n')
		# for p in self.content:
		# 	print(p.text)

			
class Webscrape():
	def __init__(self, path_storage):
		self.path_storage = path_storage

	def FormatDate(self, date):
		# print(f'Original date from link:   "{date}"')
		list_years =[]
		year1 = 2010
		while year1 < 2050:
			year1+=1
			list_years.append(year1)
		for year in list_years:
			if str(year) in date:
				# print(f'year found was "{year}" in date: "{date}"')
				idx_year = date.index(str(year))+4
				date = date[:idx_year]
				# print(f'updated string: "{date}"')
				break

		list_fullmonths = 'January February March April May June July August September October November December'.split()
		list_months = 'Jan Feb Mar Apr May Jun July Aug Sept Oct Nov Dec'.split(' ')
		# dict_months = {}
		# print(f'length of list_fullmonths:"{len(list_fullmonths)}",  length of list_months: "{len(list_months)}"')
		lowercaps = False
		lowercapsFULL = False
		uppercapsFULL = False

		# print('Testing for Case 1: Upper 1st letter Only in Full Month word')
		for idx,month in enumerate(list_fullmonths):
			# dict_months[month] = list_months[idx]
			if month in date:
				# print(f'the month "{month}" was found in the string "{date}"')
				date = date.replace(month, list_months[idx]).replace('.','')
				idx_month = date.index(list_months[idx])
				# print(f'idx_month found was "{idx_month}"')
				lowercaps = True
				lowercapsFULL = True
				uppercapsFULL = True
				break	

		if lowercaps == False:
			# print('Testing for Case 2: Upper 1st letter Only')	
			for idx,month in enumerate(list_months):
				# print(month)
				if month in date:
					# print(f'the month "{month}" was found in the string "{date}"')
					date = date.replace(month, list_months[idx]).replace('.','')
					idx_month = date.index(list_months[idx])
					# print(f'idx_month found was "{idx_month}"')
					lowercapsFULL = True
					uppercapsFULL = True
					break	

		if lowercapsFULL == False:
			# print('Testing for Case 3: Upper case for entire Full Month word')
			for idx,month in enumerate(list_fullmonths):
				# dict_months[month] = list_months[idx]
				month = month.upper()
				# print(month)
				if month in date:
					# print(f'the month "{month}" was found in the string "{date}"')
					date = date.replace(month, list_months[idx]).replace('.','')
					idx_month = date.index(list_months[idx])
					# print(f'idx_month found was "{idx_month}"')
					uppercapsFULL = True
					break	

		if uppercapsFULL == False:
			# print('Testing for Case 4: Upper case for SHORTENED Month')
			for idx,month in enumerate(list_months):
				month = month.upper()
				# print(month)
				if month in date:
					# print(f'the month "{month}" was found in the string "{date}"')
					date = date.replace(month, list_months[idx]).replace('.','')
					idx_month = date.index(list_months[idx])
					# print(f'idx_month found was "{idx_month}"')
					break	


		# print(f'idx_month is "{idx_month}"')				
		date = date[idx_month:]
		# print(f'new_date is "{date}"')

		#date should be in the format of "Month. day, year": for example == 'Oct. 2, 2020 7:27 AM ET'
		date = date.split()[:3]
		day_int = int(date[1].replace(',',''))
		if day_int<10:
			date[1]='0'+str(day_int)
		date = ' '.join(date).replace(',','')
		# print(f'date found for FormatDate(): "{date}"')
		f_article_day = datetime.datetime.strptime(date, '%b %d %Y').date()
		# print(f_article_day)
		return f_article_day
	def Article_type(self, url):  
	    list_articleTypes = ['oilprice', 'boereport', 'transcript', 'seekingalpha.com/article', 'cnbc.com', 'nbcnews', 'alberta.ca', 'seekingalpha.com/news']
	    idx=0

	    while True:
	        RegexCondition = re.search(list_articleTypes[idx],url)
	        if RegexCondition:
	            article_type=RegexCondition.group()
	            break
	        else:
	            idx+=1
	    return article_type

	def AccessWebpage(self, driver, link, article_type):
		# ['oilprice', 'boereport', 'cnbc.com', 'nbcnews', 'alberta.ca', 'seekingalpha.com/news', 'transcript', 'seekingalpha.com/article']
		driver.get(link)   

		SEL = SEL_Scrapers(driver, link)
		if article_type == 'oilprice':
			SEL.oilprice()
		elif article_type == 'boereport':
			SEL.boereport()
		elif article_type == 'cnbc.com':
			SEL.cnbc()
		elif article_type == 'nbcnews':
			SEL.nbc()
		elif article_type == 'alberta.ca':
			SEL.AB()

		SEL.Create_docx()
		SEL.Reset_PgINFO()



main()
















